# modules/sales/doc_maker.py

import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# [Helper] 테이블 셀 배경색 지정 함수
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def create_offer_sheet(form_data: dict, items: list, signature_img=None, labels: dict = None) -> io.BytesIO:
    """전문 오퍼시트 양식 생성 (다국어 지원)"""
    
    # 기본 영어 라벨
    if labels is None:
        labels = {
            "offer_sheet": "OFFER SHEET",
            "messrs": "Messrs:",
            "offer_no": "Offer No.:",
            "date": "Date:",
            "origin": "Origin",
            "shipment": "Shipment",
            "loading_port": "Loading Port",
            "destination": "Destination",
            "payment": "Payment",
            "packing": "Packing",
            "insurance": "Insurance",
            "validity": "Validity",
            "no": "No.",
            "description_of_goods": "Description of Goods",
            "quantity": "Qty",
            "unit_price": "Unit Price",
            "amount": "Amount",
            "total_amount": "TOTAL AMOUNT (FOB/CIF/CFR):",
            "dispute_resolution": "Dispute Resolution",
            "method": "Method:",
            "dispute_resolution": "Dispute Resolution",
            "governing_law": "Governing Law:",
            "accepted_by_buyer": "ACCEPTED BY (Buyer) :",
            "yours_faithfully": "Yours Faithfully,",
            "authorized_signature": "Authorized Signature"
        }
    
    doc = Document()
    
    # 1. 페이지 설정 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # 2. 헤더 (타이틀 & 판매자 정보)
    p_title = doc.add_paragraph(labels["offer_sheet"])
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].bold = True
    p_title.runs[0].font.size = Pt(20)

    p_company = doc.add_paragraph(f"[{form_data.get('seller_name', 'Company Name')}]")
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_company.runs[0].bold = True
    p_company.runs[0].font.size = Pt(14)

    p_addr = doc.add_paragraph(f"{form_data.get('seller_addr', '')}")
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr.runs[0].font.size = Pt(9)
    p_addr.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p_line = doc.add_paragraph("-" * 100)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.runs[0].font.size = Pt(8)

    # 3. 수신자 및 기본 정보 테이블
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Cm(10)
    info_table.columns[1].width = Cm(6)

    # Row 0
    cell_00 = info_table.cell(0, 0)
    cell_00.paragraphs[0].add_run(labels["messrs"] + " ").bold = True
    cell_00.paragraphs[0].add_run(form_data.get('buyer_company', ''))

    cell_01 = info_table.cell(0, 1)
    p = cell_01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(labels["offer_no"] + " ").bold = True
    p.add_run(form_data.get('offer_no', ''))

    # Row 1
    cell_10 = info_table.cell(1, 0)
    cell_10.paragraphs[0].add_run("Address: ").bold = True
    cell_10.paragraphs[0].add_run(form_data.get('address_attn', ''))

    cell_11 = info_table.cell(1, 1)
    p = cell_11.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(labels["date"] + " ").bold = True
    p.add_run(form_data.get('date', ''))

    # Row 2
    info_table.cell(2, 0).paragraphs[0].add_run("Attn: ").bold = True

    doc.add_paragraph().space_after = Pt(6)
    doc.add_paragraph("We are pleased to offer you the following goods on the terms and conditions set forth below:").space_after = Pt(8)

    # 4. 거래 조건 (Terms)
    fields = [
        (labels["origin"], "origin"),
        (labels["shipment"], "shipment"),
        (labels["loading_port"], "loading_port"),
        (labels["destination"], "destination"),
        (labels["payment"], "payment"),
        (labels["packing"], "packing"),
        (labels["insurance"], "insurance"),
        (labels["validity"], "validity")
    ]
    
    fields_table = doc.add_table(rows=len(fields), cols=2)
    for i, (label, key) in enumerate(fields):
        cell_lbl = fields_table.cell(i, 0)
        cell_lbl.text = label
        cell_lbl.paragraphs[0].runs[0].bold = True
        cell_lbl.width = Cm(4)
        fields_table.cell(i, 1).text = f": {form_data.get(key, '')}"
    
    doc.add_paragraph().space_after = Pt(8)

    # 5. 상품 테이블 (Items)
    col_headers = [
        labels["no"], 
        labels["description_of_goods"], 
        labels["quantity"], 
        labels["unit_price"], 
        labels["amount"]
    ]
    valid_items = [it for it in items if it.get("description")]
    
    goods_table = doc.add_table(rows=len(valid_items)+2, cols=5)
    goods_table.style = 'Table Grid'
    
    # 헤더 스타일링
    hdr_cells = goods_table.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "E8E8E8")
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 데이터 채우기
    for i, item in enumerate(valid_items):
        row_cells = goods_table.rows[i+1].cells
        row_cells[0].text = str(item.get("no", i+1))
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = item.get("description", "")
        row_cells[2].text = item.get("quantity", "")
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].text = item.get("unit_price", "")
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[4].text = item.get("amount", "")
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 합계 행
    total_row = goods_table.rows[-1]
    total_row.cells[0].merge(total_row.cells[3])
    total_row.cells[0].text = labels["total_amount"]
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[4].text = f"USD {form_data.get('total_amount', '')}"
    total_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph().space_after = Pt(12)

    # 6. 분쟁 해결 조항
    doc.add_paragraph("Dispute Resolution").runs[0].bold = True
    p = doc.add_paragraph()
    p.add_run("1. Method: ").bold = True

    # 7. 서명
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.cell(0, 0).text = labels["accepted_by_buyer"]
    sig_table.cell(2, 0).text = "__________________________\n" + labels["authorized_signature"]
    
    sig_table.cell(0, 1).text = labels["yours_faithfully"]
    sig_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    c_sell = sig_table.cell(1, 1)
    c_sell.text = f"[{form_data.get('seller_name', '')}]"
    c_sell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c_sell.paragraphs[0].runs[0].bold = True
    
    c_sign = sig_table.cell(2, 1)
    p_sign = c_sign.paragraphs[0]
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if signature_img:
        try:
            signature_img.seek(0)
            run = p_sign.add_run()
            run.add_picture(signature_img, width=Cm(2.5))
            p_sign.add_run("\n" + labels["authorized_signature"])
        except:
            p_sign.add_run("__________________________\n" + labels["authorized_signature"])
    else:
        p_sign.add_run("__________________________\n" + labels["authorized_signature"])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer